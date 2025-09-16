import os
import io
import logging
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile
import base64

# PDF processing
import PyPDF2
import pdfplumber
from pdf2image import convert_from_bytes

# Image processing and OCR
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import easyocr
import cv2
import numpy as np

# Document processing
import docx
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processor with OCR capabilities"""
    
    def __init__(self):
        """Initialize document processor with OCR engines"""
        self.easyocr_reader = None
        self.supported_formats = {
            'text': ['.txt', '.md'],
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']
        }
        
        # Initialize EasyOCR (lazy loading)
        try:
            self.easyocr_reader = easyocr.Reader(['en'])
            logger.info("EasyOCR initialized successfully")
        except Exception as e:
            logger.warning(f"EasyOCR initialization failed: {e}")
    
    def detect_file_type(self, filename: str, content: bytes) -> str:
        """Detect file type from filename and content"""
        file_ext = Path(filename).suffix.lower()
        
        for file_type, extensions in self.supported_formats.items():
            if file_ext in extensions:
                return file_type
        
        # Fallback: detect by content
        if content.startswith(b'%PDF'):
            return 'pdf'
        elif content.startswith(b'PK'):  # ZIP-based formats like DOCX
            return 'docx'
        elif any(content.startswith(sig) for sig in [b'\xff\xd8', b'\x89PNG', b'GIF']):
            return 'image'
        else:
            return 'text'
    
    def preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results"""
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Apply slight blur to reduce noise
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            # Convert to numpy array for OpenCV processing
            img_array = np.array(image)
            
            # Apply adaptive thresholding
            img_array = cv2.adaptiveThreshold(
                img_array, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Convert back to PIL Image
            return Image.fromarray(img_array)
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image
    
    def extract_text_with_ocr(self, image: Image.Image, use_easyocr: bool = True) -> Tuple[str, float]:
        """Extract text from image using OCR with confidence scoring"""
        text = ""
        confidence = 0.0
        
        # Preprocess image
        processed_image = self.preprocess_image(image)
        
        try:
            if use_easyocr and self.easyocr_reader:
                # Use EasyOCR (generally more accurate)
                results = self.easyocr_reader.readtext(np.array(processed_image))
                
                if results:
                    text_parts = []
                    confidences = []
                    
                    for (bbox, text_part, conf) in results:
                        text_parts.append(text_part)
                        confidences.append(conf)
                    
                    text = ' '.join(text_parts)
                    confidence = sum(confidences) / len(confidences) if confidences else 0.0
                
            else:
                # Fallback to Tesseract
                # Get text with confidence data
                data = pytesseract.image_to_data(processed_image, output_type=pytesseract.Output.DICT)
                
                text_parts = []
                confidences = []
                
                for i, conf in enumerate(data['conf']):
                    if int(conf) > 0:  # Only include confident detections
                        word = data['text'][i].strip()
                        if word:
                            text_parts.append(word)
                            confidences.append(int(conf))
                
                text = ' '.join(text_parts)
                confidence = sum(confidences) / len(confidences) / 100.0 if confidences else 0.0
                
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            # Last resort: basic Tesseract
            try:
                text = pytesseract.image_to_string(processed_image)
                confidence = 0.5  # Assume moderate confidence
            except:
                text = ""
                confidence = 0.0
        
        return text.strip(), confidence
    
    def extract_text_from_pdf(self, content: bytes) -> Dict[str, Any]:
        """Extract text from PDF with OCR fallback"""
        result = {
            'text': '',
            'method': 'unknown',
            'confidence': 1.0,
            'pages_processed': 0,
            'ocr_pages': 0
        }
        
        try:
            # First, try text extraction with pdfplumber
            with io.BytesIO(content) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    text_parts = []
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            text_parts.append(page_text)
                            result['pages_processed'] += 1
                        else:
                            # Page has no extractable text, try OCR
                            logger.info(f"Page {page_num + 1} has no text, attempting OCR...")
                            
                            try:
                                # Convert PDF page to image
                                images = convert_from_bytes(
                                    content, 
                                    first_page=page_num + 1, 
                                    last_page=page_num + 1,
                                    dpi=300
                                )
                                
                                if images:
                                    ocr_text, ocr_conf = self.extract_text_with_ocr(images[0])
                                    if ocr_text.strip():
                                        text_parts.append(ocr_text)
                                        result['ocr_pages'] += 1
                                        result['confidence'] = min(result['confidence'], ocr_conf)
                                        
                            except Exception as ocr_error:
                                logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
                    
                    result['text'] = '\n\n'.join(text_parts)
                    result['method'] = 'pdfplumber_with_ocr' if result['ocr_pages'] > 0 else 'pdfplumber'
                    result['pages_processed'] = len(pdf.pages)
                    
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2...")
            
            # Fallback to PyPDF2
            try:
                with io.BytesIO(content) as pdf_buffer:
                    pdf_reader = PyPDF2.PdfReader(pdf_buffer)
                    text_parts = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                    
                    if text_parts:
                        result['text'] = '\n\n'.join(text_parts)
                        result['method'] = 'pypdf2'
                        result['pages_processed'] = len(pdf_reader.pages)
                    else:
                        # No text found, try full OCR
                        logger.info("No text extracted with PyPDF2, attempting full OCR...")
                        result = self._full_pdf_ocr(content)
                        
            except Exception as e2:
                logger.error(f"PyPDF2 extraction failed: {e2}, attempting full OCR...")
                result = self._full_pdf_ocr(content)
        
        return result
    
    def _full_pdf_ocr(self, content: bytes) -> Dict[str, Any]:
        """Perform OCR on entire PDF when text extraction fails"""
        result = {
            'text': '',
            'method': 'full_ocr',
            'confidence': 0.0,
            'pages_processed': 0,
            'ocr_pages': 0
        }
        
        try:
            # Convert all pages to images
            images = convert_from_bytes(content, dpi=300)
            text_parts = []
            confidences = []
            
            for page_num, image in enumerate(images):
                logger.info(f"OCR processing page {page_num + 1}/{len(images)}")
                
                ocr_text, ocr_conf = self.extract_text_with_ocr(image)
                if ocr_text.strip():
                    text_parts.append(ocr_text)
                    confidences.append(ocr_conf)
                
                result['ocr_pages'] += 1
            
            result['text'] = '\n\n'.join(text_parts)
            result['confidence'] = sum(confidences) / len(confidences) if confidences else 0.0
            result['pages_processed'] = len(images)
            
        except Exception as e:
            logger.error(f"Full PDF OCR failed: {e}")
            result['text'] = "Error: Could not extract text from PDF"
            result['confidence'] = 0.0
        
        return result
    
    def extract_text_from_docx(self, content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        result = {
            'text': '',
            'method': 'docx',
            'confidence': 1.0,
            'pages_processed': 1,
            'ocr_pages': 0
        }
        
        try:
            with io.BytesIO(content) as docx_buffer:
                doc = Document(docx_buffer)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                
                result['text'] = '\n\n'.join(text_parts)
                
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            result['text'] = "Error: Could not extract text from DOCX file"
            result['confidence'] = 0.0
        
        return result
    
    def extract_text_from_image(self, content: bytes) -> Dict[str, Any]:
        """Extract text from image file using OCR"""
        result = {
            'text': '',
            'method': 'ocr',
            'confidence': 0.0,
            'pages_processed': 1,
            'ocr_pages': 1
        }
        
        try:
            image = Image.open(io.BytesIO(content))
            text, confidence = self.extract_text_with_ocr(image)
            
            result['text'] = text
            result['confidence'] = confidence
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            result['text'] = "Error: Could not extract text from image"
            result['confidence'] = 0.0
        
        return result
    
    def process_document(self, filename: str, content: bytes) -> Dict[str, Any]:
        """
        Main document processing function
        
        Args:
            filename: Name of the file
            content: File content as bytes
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_type = self.detect_file_type(filename, content)
        
        logger.info(f"Processing {filename} as {file_type} ({len(content)} bytes)")
        
        # Initialize result
        result = {
            'filename': filename,
            'file_type': file_type,
            'file_size': len(content),
            'text': '',
            'extraction_method': 'unknown',
            'confidence': 0.0,
            'pages_processed': 0,
            'ocr_pages': 0,
            'success': False,
            'error': None
        }
        
        try:
            if file_type == 'text':
                # Handle text files
                try:
                    text = content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try other encodings
                    for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            text = content.decode(encoding)
                            break
                        except:
                            continue
                    else:
                        text = content.decode('utf-8', errors='ignore')
                
                result.update({
                    'text': text,
                    'extraction_method': 'text_decode',
                    'confidence': 1.0,
                    'pages_processed': 1,
                    'success': True
                })
                
            elif file_type == 'pdf':
                # Handle PDF files
                pdf_result = self.extract_text_from_pdf(content)
                result.update(pdf_result)
                result['extraction_method'] = pdf_result['method']
                result['success'] = bool(pdf_result['text'].strip())
                
            elif file_type == 'docx':
                # Handle DOCX files
                docx_result = self.extract_text_from_docx(content)
                result.update(docx_result)
                result['extraction_method'] = docx_result['method']
                result['success'] = bool(docx_result['text'].strip())
                
            elif file_type == 'image':
                # Handle image files
                image_result = self.extract_text_from_image(content)
                result.update(image_result)
                result['extraction_method'] = image_result['method']
                result['success'] = bool(image_result['text'].strip())
                
            else:
                result['error'] = f"Unsupported file type: {file_type}"
                result['success'] = False
            
            # Validate extracted text
            if result['success'] and not result['text'].strip():
                result['success'] = False
                result['error'] = "No text could be extracted from the document"
            
            # Log processing results
            if result['success']:
                logger.info(f"Successfully processed {filename}: "
                          f"{len(result['text'])} chars, "
                          f"method: {result['extraction_method']}, "
                          f"confidence: {result['confidence']:.2f}")
            else:
                logger.warning(f"Failed to process {filename}: {result.get('error', 'Unknown error')}")
                
        except Exception as e:
            logger.error(f"Document processing error for {filename}: {str(e)}")
            result.update({
                'success': False,
                'error': str(e)
            })
        
        return result

# Global instance
document_processor = DocumentProcessor()
